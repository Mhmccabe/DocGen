# document_loader.py
import os
import pdfplumber
from docx import Document
from pathlib import Path
from zipfile import BadZipFile


def load_document(file_path):
    """
    Load and extract text from a document file.
    
    Args:
        file_path (str): Path to the document file
        
    Returns:
        str: Extracted text from the document
        
    Raises:
        FileNotFoundError: If the file doesn't exist
        ValueError: If the file type is not supported or file is corrupted
        Exception: For other processing errors
    """
    # Convert to Path object for better path handling
    file_path = Path(file_path)
    
    # Check if file exists
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Check if file is empty
    if file_path.stat().st_size == 0:
        raise ValueError(f"File is empty: {file_path}")
    
    # Get file extension
    ext = file_path.suffix.lower()
    text = ""
    
    try:
        if ext == ".pdf":
            try:
                with pdfplumber.open(file_path) as pdf:
                    # Extract text from each page
                    for page in pdf.pages:
                        # Extract text with better handling of layouts and formatting
                        page_text = page.extract_text(x_tolerance=3, y_tolerance=3)
                        if page_text:
                            text += page_text + "\n"
                        
                        # Extract tables if present and format them
                        tables = page.extract_tables()
                        for table in tables:
                            for row in table:
                                # Filter out None values and empty strings
                                row_text = [str(cell) for cell in row if cell and str(cell).strip()]
                                if row_text:
                                    text += " | ".join(row_text) + "\n"
            except Exception as e:
                raise ValueError(f"Invalid or corrupted PDF file: {str(e)}")
                
        elif ext in [".docx", ".doc"]:
            try:
                doc = Document(file_path)
                # Extract text from paragraphs
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        text += "\n" + " | ".join([cell.text for cell in row.cells])
            except BadZipFile:
                raise ValueError(f"The file appears to be corrupted or not a valid Word document. Please ensure it's a proper .docx file.")
            except Exception as e:
                raise ValueError(f"Error reading Word document: {str(e)}")
        else:
            raise ValueError(f"Unsupported file type: {ext}")
        
        # Check if any text was extracted
        if not text.strip():
            raise ValueError(f"No text could be extracted from the file. The file might be empty, corrupted, or password protected.")
            
        return text.strip()
        
    except (FileNotFoundError, ValueError):
        raise
    except Exception as e:
        raise Exception(f"Error processing {file_path}: {str(e)}")


def convert_to_markdown(text):
    """
    Convert extracted text to markdown format.
    
    Args:
        text (str): Text to convert
        
    Returns:
        str: Text in markdown format
    """
    # For demonstration, assume that plain text is acceptable as Markdown.
    # You can extend this by applying markdown formatting as needed.
    return text
