from typing import Dict, List, Optional
from pathlib import Path
import docx
from pypdf import PdfReader
from pydantic import BaseModel

class Document(BaseModel):
    """Represents a processed document with its content and metadata."""
    content: str
    file_path: str
    file_type: str
    metadata: Dict = {}

class DocumentProcessor:
    """Handles document loading and processing for different file types."""
    
    @staticmethod
    def load_document(file_path: str) -> Document:
        """
        Load a document from a file path and return a Document object.
        Supports PDF and DOCX formats.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        file_type = path.suffix.lower()
        content = ""
        
        if file_type == '.pdf':
            reader = PdfReader(file_path)
            content = "\n".join(page.extract_text() for page in reader.pages)
        elif file_type == '.docx':
            doc = docx.Document(file_path)
            content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
            
        return Document(
            content=content,
            file_path=str(path),
            file_type=file_type,
            metadata={"page_count": len(reader.pages) if file_type == '.pdf' else len(doc.paragraphs)}
        ) 